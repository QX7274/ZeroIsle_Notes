from __future__ import annotations

import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\ZeroIsle_Notes")
OUTPUT_DIR = ROOT / "output" / "结题材料_20260330"
DOCX_PATH = OUTPUT_DIR / "零屿笔记_完整结题报告_20260330.docx"
REPORT_MD_PATH = OUTPUT_DIR / "零屿笔记_完整结题报告_20260330.md"
INTRO_MD_PATH = OUTPUT_DIR / "零屿笔记_项目说明文档.md"
ARCH_MD_PATH = OUTPUT_DIR / "零屿笔记_系统架构文档.md"
TREE_MD_PATH = OUTPUT_DIR / "零屿笔记_代码目录说明.md"
MODULE_MD_PATH = OUTPUT_DIR / "零屿笔记_功能模块说明.md"
CHECKLIST_PATH = OUTPUT_DIR / "上传材料清单.txt"
ZIP_PATH = OUTPUT_DIR / "零屿笔记_核心源码与结题材料_20260330.zip"


def clean(text: str) -> str:
    return textwrap.dedent(text).strip()


REPORT_TITLE = "《零屿笔记》创新型智能笔记应用结题报告"


REPORT_SECTIONS = [
    (
        "一、项目基本信息",
        [
            clean(
                """
                项目名称：零屿笔记创新型智能笔记应用
                项目类别：大学生创新创业训练项目
                项目负责人：钱鑫
                指导教师：王丽娟
                结题整理日期：2026年3月30日

                说明：本报告根据项目申报书《零屿笔记创新型智能笔记应用》以及项目当前代码仓库、说明文档、测试文件和部署配置综合整理形成，
                重点突出项目执行情况、研究总结、软件系统完整性及上传附件说明。经费部分按项目实际情况填写为“未使用经费，无报销”。
                """
            )
        ],
    ),
    (
        "二、项目执行情况",
        [
            "（一）本项目的选题背景、目的与意义",
            clean(
                """
                随着高校学生在学习过程中产生的课堂笔记、录音资料、图片截图、课件文档等多模态资料不断增多，传统笔记工具虽然能够完成基础记录，
                但普遍存在功能碎片化、智能辅助不足、知识关联能力弱、复习提醒支持不够等问题。尤其在教育学习场景中，学生往往需要在手写笔记、
                语音记录、PPT截图、教材摘录和网络资料之间频繁切换，后续整理和复习成本较高，难以形成高质量的个人知识体系。

                基于上述背景，本项目提出“零屿笔记”智能笔记应用，目标是面向教育与学习场景，构建一套融合笔记编辑、语音转写、AI文本辅助、
                知识图谱、提醒管理、文件处理、社区交流等功能于一体的综合性智能学习工具。项目希望借助人工智能技术与移动端开发能力，
                将笔记工具由“被动记录”进一步升级为“主动整理、智能分析、关联检索和学习支持”的知识管理平台。

                本项目具有较明显的实践意义和应用价值。一方面，项目尝试将 OCR、语音识别、知识图谱、文本生成与多模态交互能力应用到学习场景中，
                推动 AI 技术与教育工具融合；另一方面，项目成果面向学生实际需求，能够在课堂记录、课后整理、知识复习、资料检索与学习交流等方面
                提供帮助，具备较好的校园应用前景和后续推广潜力。
                """
            ),
            "（二）项目的创新点与特色",
            clean(
                """
                1. 多功能一体化集成。项目并非单一笔记应用，而是围绕学习场景将笔记记录、AI处理、语音识别、知识组织、提醒复习、社区共享等功能
                统一到同一系统中，形成较完整的软件功能闭环。

                2. 多模态交互能力。项目支持文本、语音、图片、文档等多种信息形态的处理与输入，并结合搜索、摘要、翻译、图像分析、语音转写等能力，
                提升学习资料数字化和再利用效率。

                3. 面向知识管理的组织方式。项目引入知识图谱、分类标签、语义搜索、内容关联等思路，帮助用户从零散笔记逐步构建结构化个人知识体系，
                强化“知识连接创造价值”的产品理念。

                4. 跨平台与工程化实现。项目采用 React Native 开发移动端，配套 Django 后端、MongoDB/Realm 数据存储、Neo4j 图谱支撑、
                Redis/Celery 异步能力及测试与部署文档，已形成具有较强完整性的软件系统，而非停留在概念原型阶段。

                5. 软件完整性较强。当前仓库已形成移动端、后端、管理后台、文档体系、测试用例、部署配置等内容，能够较好展示项目从立项、设计、
                开发到阶段验证的完整实践过程。
                """
            ),
            "（三）团队成员分工和合作情况",
            clean(
                """
                项目实施过程中，团队按照“项目统筹、前端开发、后端开发、算法支持、测试验证”进行协同分工。

                1. 钱鑫：负责项目整体统筹与推进，承担核心功能整合、知识管理相关模块推进、软件系统集成、项目材料整理与结题汇报工作。
                2. 王钰榕：主要参与前端页面开发、交互设计、移动端界面联调和用户体验优化。
                3. 王荣琪：主要参与 OCR 与识别相关功能研究、识别能力优化及相关技术支持。
                4. 张明峰：主要参与后端服务开发、接口设计、数据管理和系统支撑工作。
                5. 黄晋豪：主要参与语音相关功能、功能验证与测试支持工作。

                指导教师王丽娟老师在项目选题、技术路线把控、功能推进、实施节奏安排及结题总结方面给予了持续指导。团队在项目推进过程中能够依据开发阶段
                及时调整协作重点，通过文档共享、模块拆分、接口联调和阶段汇总等方式保持较好的合作效率。
                """
            ),
            "（四）成果简述",
            clean(
                """
                项目目前已经形成较为完整的软件成果，完成了“零屿笔记”智能笔记系统的主体开发工作。仓库中已包含 React Native 移动端代码、
                Django 后端代码、Android/iOS 原生模块、管理后台、部署配置、说明文档及基础测试文件。系统已实现或基本实现的主要功能包括：
                用户注册登录、笔记创建与编辑、分类标签管理、文件上传与处理、提醒管理、搜索检索、AI 文本辅助、语音转文字、知识关联、社区互动等。

                从工程成果上看，项目已具备 README、架构文档、安装部署文档、功能说明文档、模块状态说明、测试配置等材料。项目不是仅有界面展示，
                而是形成了包含前端交互、后端服务、数据管理、接口能力和文档体系在内的较完整软件系统，已具备演示、验证和持续优化的基础。
                """
            ),
        ],
    ),
    (
        "三、研究总结报告",
        [
            "（一）预定计划执行情况、项目研究和实践情况（含进度安排、完成内容、关键技术及效果等）",
            clean(
                """
                项目总体上按照申报书的技术路线与阶段安排推进。前期主要完成了需求分析、架构设计、技术选型、前后端基础框架搭建以及移动端基本页面结构建设；
                中期围绕笔记系统、AI助手、搜索、提醒、语音与文件处理等核心功能持续开发与联调；后期则进一步补充知识管理、社区协作、部署文档、
                测试文件和系统整合工作，逐步形成了较完整的项目成果。

                在完成内容方面，前端已经覆盖笔记、AI、搜索、知识管理、提醒、社区等主要页面；后端已形成 users、notes、reminder、search、
                ai_assistant、voice_recognition、knowledge_graph、community、sync 等功能模块；数据层面结合 Realm、本地存储、MongoDB、Neo4j
                等方案进行支撑；工程化方面具备 package.json、Docker 配置、测试脚本、部署文档与模块说明。根据 README 与现有目录结构，
                项目当前版本可归纳为“核心功能完备、系统结构完整、部分高级功能持续优化”的状态。

                关键技术主要包括：React Native 跨平台开发、Redux 状态管理、Django REST Framework 接口服务、MongoDB/Realm 数据存储、
                Neo4j 图数据库、Whisper 语音识别、OCR 与图像处理、AI 文本处理、多模态搜索、同步与缓存机制等。这些技术共同支撑了本项目从方案设想到软件落地。

                从效果上看，项目已完成一套可演示、可阅读、可继续完善的智能笔记系统。仓库中存在清晰的功能模块划分、较为丰富的说明文档与测试文件，
                并能够通过最小前后端测试验证测试框架正常运行。2026年3月30日已在本地执行前端最小 Jest 测试和后端最小 Pytest 测试，结果均通过，
                说明项目已具备基本的工程化验证条件。
                """
            ),
            "（二）项目实施的收获与体会、项目工作有哪些不足、项目工作中困难与解决方法",
            clean(
                """
                通过本项目的实施，团队对完整软件项目的推进流程有了更加深入的理解，尤其在需求分析、功能拆分、技术选型、跨平台开发、接口联调、
                文档整理、测试验证和项目总结等方面积累了较为系统的实践经验。团队成员在项目过程中进一步提升了协作能力与工程意识，
                对如何将人工智能技术与教育学习场景结合也有了更加明确的认识。

                项目实施过程中也存在一定不足。首先，由于项目功能覆盖面较广，部分高级功能仍处于持续优化阶段，少量模块在不同平台上的体验一致性仍需加强。
                其次，团队在有限周期内优先保障了系统主体开发和核心功能落地，对于大规模用户试点、论文专利产出、深度算法优化等工作推进相对有限。
                再次，虽然已建立基础测试体系，但系统级自动化测试、性能测试和更大规模的真实使用验证仍有提升空间。

                项目推进中遇到的主要困难包括：模块多、集成复杂、AI 与识别类功能调试成本较高、前后端联调工作量较大。为解决这些问题，团队采取了分阶段迭代、
                模块化开发、优先核心功能、依托文档沉淀推进协作、逐步补充测试与部署说明等方法，最终保证了项目能够形成结构完整、可展示、可继续迭代的成果。
                这些经历也让团队更加明确，创新创业项目不仅要有想法，还要重视软件工程落地、文档规范和阶段总结。
                """
            ),
        ],
    ),
    (
        "四、经费开支与报销情况",
        [
            clean(
                """
                本项目实施期间未实际使用学校拨付经费或其他配套经费，项目开发、测试与材料整理主要依靠团队现有设备、已有开发环境及校内实验条件完成。
                因此，本项目经费开支与报销情况可填写为：

                未使用经费，无报销。
                """
            )
        ],
    ),
    (
        "五、项目说明文档",
        [
            clean(
                """
                零屿笔记是一款围绕学习与知识管理场景开发的智能笔记应用，主要面向课堂记录、课后整理、资料归档、知识关联和学习复习等需求。
                项目在移动端提供笔记、AI、搜索、提醒、知识管理、社区互动等核心页面，并通过后端服务支撑用户认证、数据管理、搜索分析、
                语音处理和知识图谱等能力。整体设计强调“记录、整理、关联、复习、交流”的连续使用链路。

                从软件形态上看，项目已经形成较完整的工程结构，包含移动端前端、后端接口服务、数据库与本地存储层、原生模块扩展、
                文档说明与测试配置。现有 README、安装部署指南、功能说明、架构文档等材料能够支撑项目展示、验收说明和后续迭代维护。
                """
            )
        ],
    ),
    (
        "六、系统架构文档",
        [
            clean(
                """
                1. 前端层：基于 React Native 构建移动应用，包含页面、组件、导航、状态管理、服务调用和部分原生能力桥接。
                2. 原生扩展层：Android 与 iOS 目录下实现了语音、通知、文件、画布、PDF、分页笔记等原生模块，用于增强跨平台体验与性能。
                3. 后端服务层：基于 Django 和 Django REST Framework 搭建，提供用户、笔记、提醒、搜索、AI助手、知识图谱、社区、同步等业务接口。
                4. 数据层：结合 Realm、本地存储、MongoDB、Neo4j、Redis 等方案，分别处理本地离线数据、在线业务数据、图谱关系和缓存/任务支撑。
                5. AI与处理层：接入文本处理、语音识别、图像分析、OCR、搜索与推荐能力，为笔记内容提供智能增强。

                整体架构采用前后端分离思路，强调模块解耦与多端协同。前端负责交互与本地能力承载，后端负责业务处理与接口统一，
                数据层负责信息持久化和知识组织，文档与测试体系则为项目交付和维护提供支撑。
                """
            )
        ],
    ),
    (
        "七、代码目录说明",
        [
            clean(
                """
                1. src/：React Native 前端主体代码，包含 components、screens、services、redux、navigation、hooks、models 等目录。
                2. backend/：Django 后端主体代码，包含 users、notes、search、reminder、ai_assistant、voice_recognition、knowledge_graph、
                   community、sync 等模块及全局配置。
                3. android/、ios/：移动端原生工程和桥接模块，用于增强文件处理、语音、通知、画布、PDF 等能力。
                4. admin_system/：项目管理后台相关代码。
                5. web/：项目网页或 Web 端相关内容。
                6. docs/、Info/：项目说明、架构、功能、部署、优化记录等文档材料。
                7. e2e/、src/tests/、backend/tests/：自动化测试、接口测试和基础测试文件。
                8. scripts/：环境检查、部署辅助和材料生成等脚本。

                目录结构总体较清晰，能够体现项目从业务功能、工程支撑到文档沉淀的完整性。
                """
            )
        ],
    ),
    (
        "八、功能模块说明",
        [
            clean(
                """
                1. 笔记管理模块：支持笔记创建、编辑、分类、标签、版本、附件、分享等基础能力。
                2. AI助手模块：支持文本摘要、改写、翻译、解释、关键词提取、语法检查等智能文本处理功能。
                3. 搜索模块：支持关键词搜索、语义搜索、语音搜索、图像搜索和搜索历史记录。
                4. 语音模块：支持录音、语音转文字、会议纪要生成和相关语音交互能力。
                5. 知识管理模块：支持知识图谱、知识库、关联推荐、结构分析等能力。
                6. 画布与思维导图模块：支持无限画布、分页笔记、思维导图和相关内容组织方式。
                7. 提醒与日程模块：支持提醒创建、分类、日历查看、统计与通知能力。
                8. 文件与文档处理模块：支持上传、预览、导出、OCR、转换等功能。
                9. 社区与协作模块：支持帖子、评论、点赞、关注、分享、群组等功能。
                10. 同步与数据管理模块：支持本地存储、离线使用、同步策略和后端数据协同。

                各模块共同构成了本项目的软件主体，使系统具备较强的完整性、可演示性和持续扩展能力。
                """
            )
        ],
    ),
    (
        "九、结题附件建议",
        [
            clean(
                """
                结题上传时建议至少提交以下材料：
                1. 本完整结题报告（DOCX）。
                2. 项目申报书 PDF。
                3. 项目说明文档。
                4. 系统架构文档。
                5. 代码目录说明。
                6. 功能模块说明。
                7. 核心源码与结题材料压缩包。

                若学校系统允许，可同时补充界面截图、答辩 PPT、功能演示视频或展板材料，以增强结题展示效果。
                """
            )
        ],
    ),
]


INTRO_MD = clean(
    """
    # 零屿笔记项目说明文档

    ## 1. 项目概述
    零屿笔记是一款面向学习与知识管理场景的智能笔记应用，围绕课堂记录、资料整理、知识关联、内容复习和交流分享等需求展开设计。
    项目整合了笔记编辑、AI 文本辅助、语音转文字、知识图谱、搜索、提醒、文件处理与社区互动等功能，目标是帮助用户从分散资料中
    逐步建立结构化的个人知识体系。

    ## 2. 软件形态
    项目不是单一原型页面，而是已经形成了包含移动端、后端服务、原生模块、文档材料和测试基础在内的较完整工程系统。
    移动端采用 React Native 开发，后端采用 Django 与 Django REST Framework，数据层结合 Realm、MongoDB、Neo4j、Redis 等方案，
    支撑离线使用、知识关系组织和接口服务。

    ## 3. 主要技术
    - 前端：React Native、Redux Toolkit、React Navigation、React Native Paper
    - 后端：Django、DRF、Celery、Channels
    - 数据：Realm、MongoDB、Neo4j、Redis
    - 智能能力：AI 文本处理、OCR、语音识别、图像分析、语义搜索

    ## 4. 已形成成果
    当前仓库已经包含前端代码、后端模块、移动端原生工程、说明文档、安装部署指南、模块状态文档和测试文件。
    系统主体功能已具备较好的展示与验证基础，能够较完整体现项目从构想到软件落地的实施过程。

    ## 5. 结题说明
    本次结题整理重点突出项目的实际完成情况、软件完整性、功能体系、目录结构和上传附件准备情况，便于学校系统直接提交。
    """
)


ARCH_MD = clean(
    """
    # 零屿笔记系统架构文档

    ## 1. 总体架构
    零屿笔记整体采用前后端分离与多层协同的架构设计：

    - 展示与交互层：React Native 移动端页面、组件、导航、状态管理
    - 原生能力层：Android/iOS 原生模块，负责语音、文件、通知、PDF、画布等能力
    - 业务服务层：Django + DRF 提供统一 API 接口
    - 数据存储层：Realm、本地缓存、MongoDB、Neo4j、Redis
    - 智能处理层：文本处理、OCR、语音识别、图像分析、知识关联与搜索

    ## 2. 前端架构
    `src/` 目录承担前端主体职责，主要由 `screens/` 页面层、`components/` 组件层、`services/` 服务层、
    `redux/` 状态管理层、`navigation/` 导航层以及 `hooks/` 和 `models/` 支撑模块组成。

    ## 3. 后端架构
    `backend/` 目录承担接口与业务处理职责，围绕用户、笔记、搜索、提醒、AI 助手、语音识别、知识图谱、社区和同步等模块组织，
    通过统一路由与服务层实现数据访问、任务分发和接口响应。

    ## 4. 数据与同步
    项目强调本地使用与云端扩展并行。移动端结合 Realm 和本地缓存支撑离线能力，后端使用 MongoDB 管理业务数据，
    Neo4j 用于关系型知识组织，Redis 与 Celery 支撑缓存和异步任务。

    ## 5. 工程与交付
    项目同时具备 README、安装部署指南、架构说明、功能说明、模块状态说明和测试文件，能够支撑开发、验收与后续迭代。
    """
)


TREE_MD = clean(
    """
    # 零屿笔记代码目录说明

    ## 1. 根目录关键内容
    - `src/`：前端核心代码
    - `backend/`：后端核心代码
    - `android/`：Android 原生工程
    - `ios/`：iOS 原生工程
    - `admin_system/`：管理后台相关代码
    - `web/`：Web 相关代码
    - `Info/`、`docs/`：项目文档
    - `scripts/`：辅助脚本

    ## 2. 前端目录说明
    - `src/components/`：通用与业务组件
    - `src/screens/`：页面级模块
    - `src/services/`：服务与接口调用
    - `src/redux/`：状态管理
    - `src/navigation/`：导航结构
    - `src/hooks/`：自定义 Hook
    - `src/models/`、`src/schemas/`：数据模型和结构定义
    - `src/tests/`：前端测试

    ## 3. 后端目录说明
    - `backend/users/`：用户认证与用户管理
    - `backend/notes/`：笔记核心业务
    - `backend/search/`：搜索与检索
    - `backend/reminder/`：提醒与日程
    - `backend/ai_assistant/`：AI 文本与图像处理
    - `backend/voice_recognition/`：语音转写与会议纪要
    - `backend/knowledge_graph/`：知识图谱与知识关联
    - `backend/community/`：社区互动
    - `backend/sync/`：同步机制
    - `backend/tests/`：后端测试

    ## 4. 目录特点
    项目目录覆盖了应用展示层、业务层、数据层、原生层、文档层和测试层，体现出较清晰的软件工程结构。
    """
)


MODULE_MD = clean(
    """
    # 零屿笔记功能模块说明

    ## 1. 笔记管理模块
    支持笔记创建、编辑、分类、标签、版本记录、附件处理和分享功能，是整个系统的核心业务模块。

    ## 2. AI 助手模块
    支持摘要、翻译、改写、解释、关键词提取、语法检查等文本处理能力，为笔记内容提供智能增强。

    ## 3. 搜索模块
    支持关键词搜索、语义搜索、语音搜索、图像搜索和搜索建议，提升资料检索效率。

    ## 4. 语音模块
    支持录音、语音转文字、会议纪要生成和相关语音处理能力，用于课堂、会议和日常记录场景。

    ## 5. 知识管理模块
    包含知识图谱、知识库、结构分析、关联推荐等能力，用于组织和关联用户知识内容。

    ## 6. 文件与文档模块
    支持文件上传、预览、导出、格式转换、OCR 和相关文档处理能力。

    ## 7. 提醒与日程模块
    支持提醒创建、任务分类、统计分析、日历视图和通知能力。

    ## 8. 画布与思维导图模块
    支持无限画布、分页笔记、思维导图和内容组织等创意型功能。

    ## 9. 社区与协作模块
    支持帖子、评论、点赞、关注、分享、群组和内容交流，增强知识共享与互动。

    ## 10. 同步与数据管理模块
    支持本地存储、离线工作、同步服务和数据管理，为系统稳定运行提供底层支撑。
    """
)


ROOT_FILES = [
    "README.md",
    "package.json",
    "package-lock.json",
    "yarn.lock",
    "index.js",
    "app.json",
    "babel.config.js",
    "metro.config.js",
    "jest.config.js",
    "tsconfig.json",
    "Dockerfile",
    "docker-compose.yml",
    "docker-compose.prod.yml",
    "electron-builder.yml",
    "requirements.txt",
    "pytest.ini",
    ".env.example",
    ".eslintrc.js",
    ".eslintignore",
    ".prettierrc.js",
]

ROOT_DIRS = [
    "src",
    "backend",
    "android",
    "ios",
    "web",
    "Info",
    "docs",
    "admin_system",
    "electron",
    "e2e",
    "scripts",
]

EXCLUDED_DIR_NAMES = {
    "node_modules",
    ".git",
    ".venv",
    "__pycache__",
    ".pytest_cache",
    ".idea",
    ".cursor",
    ".qoder",
    ".bundle",
    ".vscode",
    ".gradle",
    "backups",
    "archive",
    "build",
    "Pods",
    "DerivedData",
    "zeroislenotes_local.realm.management",
}

EXCLUDED_SUFFIXES = {
    ".pyc",
    ".pyo",
    ".log",
    ".tmp",
    ".lock",
    ".apk",
    ".aab",
    ".hprof",
}

EXCLUDED_FILE_NAMES = {
    ".env",
    "zeroislenotes_local.realm",
    "zeroislenotes_local.realm.lock",
}


def write_text(path: Path, content: str) -> None:
    path.write_text(content + "\n", encoding="utf-8")


def build_full_report_markdown() -> str:
    lines: list[str] = [f"# {REPORT_TITLE}", ""]
    for heading, parts in REPORT_SECTIONS:
        lines.append(f"## {heading}")
        lines.append("")
        for part in parts:
            if part.startswith("（"):
                lines.append(f"### {part}")
                lines.append("")
            else:
                lines.append(part)
                lines.append("")
    return "\n".join(lines).strip() + "\n"


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraphs(doc: Document, text: str) -> None:
    for para in text.split("\n\n"):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)


def build_docx() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = 1
    run = title.add_run(REPORT_TITLE)
    run.font.size = Pt(20)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    subtitle_run = subtitle.add_run("基于项目申报书与当前软件成果整理")
    subtitle_run.font.name = "宋体"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    subtitle_run.font.size = Pt(12)

    doc.add_paragraph("")

    for heading, parts in REPORT_SECTIONS:
        doc.add_heading(heading, level=1)
        for part in parts:
            if part.startswith("（"):
                doc.add_heading(part, level=2)
            else:
                add_paragraphs(doc, part)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("附录一：项目说明文档", level=1)
    add_paragraphs(doc, INTRO_MD)
    doc.add_heading("附录二：系统架构文档", level=1)
    add_paragraphs(doc, ARCH_MD)
    doc.add_heading("附录三：代码目录说明", level=1)
    add_paragraphs(doc, TREE_MD)
    doc.add_heading("附录四：功能模块说明", level=1)
    add_paragraphs(doc, MODULE_MD)

    doc.save(DOCX_PATH)


def should_include(path: Path) -> bool:
    if path.name in EXCLUDED_FILE_NAMES:
        return False
    if any(part in EXCLUDED_DIR_NAMES for part in path.parts):
        return False
    if path.suffix.lower() in EXCLUDED_SUFFIXES:
        return False
    return True


def iter_included_files(base: Path):
    for path in base.rglob("*"):
        if path.is_file() and should_include(path):
            yield path


def build_zip() -> tuple[int, int]:
    count = 0
    total_size = 0

    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
        for file_name in ROOT_FILES:
            path = ROOT / file_name
            if path.exists() and path.is_file() and should_include(path):
                zf.write(path, arcname=file_name)
                count += 1
                total_size += path.stat().st_size

        for dir_name in ROOT_DIRS:
            base = ROOT / dir_name
            if not base.exists():
                continue
            for path in iter_included_files(base):
                arcname = path.relative_to(ROOT)
                zf.write(path, arcname=str(arcname))
                count += 1
                total_size += path.stat().st_size

        generated_files = [
            DOCX_PATH,
            REPORT_MD_PATH,
            INTRO_MD_PATH,
            ARCH_MD_PATH,
            TREE_MD_PATH,
            MODULE_MD_PATH,
            CHECKLIST_PATH,
        ]
        for path in generated_files:
            if path.exists():
                arcname = Path("结题材料") / path.name
                zf.write(path, arcname=str(arcname))
                count += 1
                total_size += path.stat().st_size

    return count, total_size


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    report_md = build_full_report_markdown()
    write_text(REPORT_MD_PATH, report_md)
    write_text(INTRO_MD_PATH, INTRO_MD)
    write_text(ARCH_MD_PATH, ARCH_MD)
    write_text(TREE_MD_PATH, TREE_MD)
    write_text(MODULE_MD_PATH, MODULE_MD)

    build_docx()

    checklist = clean(
        f"""
        结题材料清单
        1. {DOCX_PATH.name}
        2. {REPORT_MD_PATH.name}
        3. {INTRO_MD_PATH.name}
        4. {ARCH_MD_PATH.name}
        5. {TREE_MD_PATH.name}
        6. {MODULE_MD_PATH.name}
        7. {ZIP_PATH.name}

        备注
        - 经费情况：未使用经费，无报销。
        - 源码压缩包仅保留核心源码、说明文档、关键配置和结题材料，
          已排除 node_modules、缓存、数据库临时文件、.env 等不适合上传的内容。
        """
    )
    write_text(CHECKLIST_PATH, checklist)

    file_count, total_size = build_zip()
    size_mb = ZIP_PATH.stat().st_size / (1024 * 1024)

    print(f"OUTPUT_DIR={OUTPUT_DIR}")
    print(f"DOCX={DOCX_PATH}")
    print(f"ZIP={ZIP_PATH}")
    print(f"ZIP_FILE_COUNT={file_count}")
    print(f"ZIP_SOURCE_SIZE_BYTES={total_size}")
    print(f"ZIP_SIZE_MB={size_mb:.2f}")


if __name__ == "__main__":
    main()
