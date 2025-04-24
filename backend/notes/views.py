"""笔记视图"""

from rest_framework import viewsets, permissions, status, filters
from rest_framework.decorators import action
from rest_framework.response import Response
from django_filters.rest_framework import DjangoFilterBackend
from django.db.models import Q, Count, F, ExpressionWrapper, FloatField
from django.db.models.functions import TruncDate, TruncMonth
from django.shortcuts import get_object_or_404
from django.contrib.auth.models import User
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.utils import timezone
from django.http import HttpResponse
from django.db import transaction
from django.conf import settings
import json
import csv
import markdown
from docx import Document
import os
import zipfile
import io
from datetime import timedelta
import difflib
import torch

from .models import Note, Category, Tag, NoteShare, NoteVersion, NoteAttachment, NoteSync, NoteComment, NoteCollaboration, NoteTemplate, NoteBackup, NoteReminder, Notification, Handwriting, HandwritingShare, OCRModel, OCRTrainingData, WhisperModel, WhisperTrainingData
from .serializers import (
    NoteListSerializer, 
    NoteDetailSerializer, 
    NoteCreateUpdateSerializer,
    CategorySerializer,
    TagSerializer,
    NoteSerializer,
    NoteShareSerializer,
    NoteShareCreateSerializer,
    NoteVersionSerializer,
    NoteAttachmentSerializer,
    NoteSyncSerializer,
    NoteCommentSerializer,
    NoteCollaborationSerializer,
    NoteTemplateSerializer,
    NoteBackupSerializer,
    NoteReminderSerializer,
    NotificationSerializer,
    HandwritingSerializer,
    HandwritingShareSerializer,
    OCRModelSerializer,
    OCRTrainingDataSerializer,
    WhisperModelSerializer,
    WhisperTrainingDataSerializer
)


class NoteViewSet(viewsets.ModelViewSet):
    """笔记视图集"""
    serializer_class = NoteSerializer
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_fields = ['is_pinned', 'is_archived']
    search_fields = ['title', 'content']
    ordering_fields = ['created_at', 'updated_at', 'title']
    
    def get_queryset(self):
        return Note.objects.filter(user=self.request.user)
    
    def get_serializer_class(self):
        if self.action == 'list':
            return NoteListSerializer
        elif self.action in ['create', 'update', 'partial_update']:
            return NoteCreateUpdateSerializer
        return NoteDetailSerializer
    
    def perform_create(self, serializer):
        serializer.save(user=self.request.user)
    
    def perform_update(self, serializer):
        instance = serializer.instance
        # 保存当前版本到历史记录
        NoteVersion.objects.create(
            note=instance,
            title=instance.title,
            content=instance.content,
            user=self.request.user
        )
        serializer.save()
    
    @action(detail=False, methods=['get'])
    def favorites(self, request):
        queryset = self.get_queryset().filter(
            is_favorite=True,
            user=request.user
        ).order_by('-updated_at')
        
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def trash(self, request):
        queryset = self.get_queryset().filter(
            is_deleted=True,
            user=request.user
        ).order_by('-deleted_at')
        
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def move_to_trash(self, request, pk=None):
        note = self.get_object()
        note.is_deleted = True
        note.deleted_at = timezone.now()
        note.save()
        return Response({'status': 'success'})
    
    @action(detail=True, methods=['post'])
    def restore(self, request, pk=None):
        note = self.get_object()
        note.is_deleted = False
        note.deleted_at = None
        note.save()
        return Response({'status': 'success'})
    
    @action(detail=True, methods=['post'])
    def toggle_favorite(self, request, pk=None):
        note = self.get_object()
        note.is_favorite = not note.is_favorite
        note.save()
        return Response({
            'status': 'success', 
            'is_favorite': note.is_favorite
        })
    
    @action(detail=False, methods=['get'])
    def search(self, request):
        query = request.query_params.get('q', '')
        category_id = request.query_params.get('category')
        tag_ids = request.query_params.getlist('tags')
        is_pinned = request.query_params.get('is_pinned')
        is_archived = request.query_params.get('is_archived')
        is_favorite = request.query_params.get('is_favorite')
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')

        queryset = self.get_queryset()

        # 搜索关键词
        if query:
            queryset = queryset.filter(
                Q(title__icontains=query) | 
                Q(content__icontains=query)
            )

        # 分类过滤
        if category_id:
            queryset = queryset.filter(category_id=category_id)

        # 标签过滤
        if tag_ids:
            queryset = queryset.filter(tags__id__in=tag_ids)

        # 状态过滤
        if is_pinned is not None:
            queryset = queryset.filter(is_pinned=is_pinned.lower() == 'true')
        if is_archived is not None:
            queryset = queryset.filter(is_archived=is_archived.lower() == 'true')
        if is_favorite is not None:
            queryset = queryset.filter(is_favorite=is_favorite.lower() == 'true')

        # 日期范围过滤
        if start_date and end_date:
            queryset = queryset.filter(
                created_at__range=[start_date, end_date]
            )

        # 排序
        sort_by = request.query_params.get('sort_by', '-updated_at')
        queryset = queryset.order_by(sort_by)

        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def statistics(self, request):
        total_notes = self.get_queryset().count()
        pinned_notes = self.get_queryset().filter(is_pinned=True).count()
        archived_notes = self.get_queryset().filter(is_archived=True).count()
        favorite_notes = self.get_queryset().filter(is_favorite=True).count()
        
        # 按分类统计
        category_stats = []
        for category in Category.objects.filter(user=request.user):
            count = self.get_queryset().filter(category=category).count()
            category_stats.append({
                'category_id': category.id,
                'category_name': category.name,
                'count': count
            })

        # 按标签统计
        tag_stats = []
        for tag in Tag.objects.filter(user=request.user):
            count = self.get_queryset().filter(tags=tag).count()
            tag_stats.append({
                'tag_id': tag.id,
                'tag_name': tag.name,
                'count': count
            })

        return Response({
            'total_notes': total_notes,
            'pinned_notes': pinned_notes,
            'archived_notes': archived_notes,
            'favorite_notes': favorite_notes,
            'category_stats': category_stats,
            'tag_stats': tag_stats
        })

    @action(detail=True, methods=['post'])
    def pin(self, request, pk=None):
        note = self.get_object()
        note.is_pinned = not note.is_pinned
        note.save()
        return Response({'status': 'success'})

    @action(detail=True, methods=['post'])
    def archive(self, request, pk=None):
        note = self.get_object()
        note.is_archived = not note.is_archived
        note.save()
        return Response({'status': 'success'})

    @action(detail=True, methods=['post'])
    def share(self, request, pk=None):
        """分享笔记"""
        note = self.get_object()
        shared_with_id = request.data.get('shared_with')
        permission = request.data.get('permission', 'view')  # view, edit, comment
        expires_at = request.data.get('expires_at')
        
        if not shared_with_id:
            return Response(
                {'error': 'shared_with is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            shared_with = User.objects.get(id=shared_with_id)
            
            # 检查是否已经分享过
            if NoteShare.objects.filter(note=note, shared_with=shared_with).exists():
                return Response(
                    {'error': 'Note is already shared with this user'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # 创建分享记录
            share = NoteShare.objects.create(
                note=note,
                shared_by=request.user,
                shared_with=shared_with,
                permission=permission,
                expires_at=expires_at
            )
            
            # 创建通知
            Notification.objects.create(
                user=shared_with,
                title='笔记已分享',
                message=f'用户 {request.user.username} 与您分享了笔记 "{note.title}"',
                type='note_shared',
                related_note=note
            )
            
            serializer = NoteShareSerializer(share)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        except User.DoesNotExist:
            return Response(
                {'error': 'User not found'},
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['get'])
    def shares(self, request, pk=None):
        """获取笔记的分享列表"""
        note = self.get_object()
        shares = note.shares.all()
        serializer = NoteShareSerializer(shares, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def update_share(self, request, pk=None):
        """更新分享权限"""
        note = self.get_object()
        share_id = request.data.get('share_id')
        permission = request.data.get('permission')
        expires_at = request.data.get('expires_at')
        
        if not share_id:
            return Response(
                {'error': 'share_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            share = note.shares.get(id=share_id)
            
            if permission:
                share.permission = permission
            if expires_at:
                share.expires_at = expires_at
            
            share.save()
            
            # 创建通知
            Notification.objects.create(
                user=share.shared_with,
                title='分享权限已更新',
                message=f'笔记 "{note.title}" 的分享权限已更新',
                type='share_updated',
                related_note=note
            )
            
            serializer = NoteShareSerializer(share)
            return Response(serializer.data)
            
        except NoteShare.DoesNotExist:
            return Response(
                {'error': 'Share not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['post'])
    def revoke_share(self, request, pk=None):
        """撤销分享"""
        note = self.get_object()
        share_id = request.data.get('share_id')
        
        if not share_id:
            return Response(
                {'error': 'share_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            share = note.shares.get(id=share_id)
            shared_with = share.shared_with
            share.delete()
            
            # 创建通知
            Notification.objects.create(
                user=shared_with,
                title='分享已撤销',
                message=f'笔记 "{note.title}" 的分享已撤销',
                type='share_revoked',
                related_note=note
            )
            
            return Response({'status': 'success'})
            
        except NoteShare.DoesNotExist:
            return Response(
                {'error': 'Share not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['post'])
    def share_link(self, request, pk=None):
        """生成分享链接"""
        note = self.get_object()
        permission = request.data.get('permission', 'view')
        expires_at = request.data.get('expires_at')
        
        # 创建分享记录
        share = NoteShare.objects.create(
            note=note,
            shared_by=request.user,
            permission=permission,
            expires_at=expires_at,
            is_link=True
        )
        
        # 生成分享链接
        share_link = f"{settings.SITE_URL}/notes/{note.id}/share/{share.id}"
        
        return Response({
            'share_link': share_link,
            'expires_at': expires_at
        })

    @action(detail=True, methods=['get'])
    def versions(self, request, pk=None):
        note = self.get_object()
        versions = note.versions.all().order_by('-created_at')
        serializer = NoteVersionSerializer(versions, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['get'])
    def compare_versions(self, request, pk=None):
        """比较笔记的两个版本"""
        note = self.get_object()
        version1_id = request.query_params.get('version1')
        version2_id = request.query_params.get('version2')
        
        if not version1_id or not version2_id:
            return Response(
                {'error': 'version1 and version2 are required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            version1 = note.versions.get(id=version1_id)
            version2 = note.versions.get(id=version2_id)
            
            # 生成差异
            differ = difflib.Differ()
            diff = list(differ.compare(
                version1.content.splitlines(),
                version2.content.splitlines()
            ))
            
            # 格式化差异
            formatted_diff = []
            for line in diff:
                if line.startswith('+ '):
                    formatted_diff.append({
                        'type': 'addition',
                        'content': line[2:]
                    })
                elif line.startswith('- '):
                    formatted_diff.append({
                        'type': 'deletion',
                        'content': line[2:]
                    })
                elif line.startswith('? '):
                    continue
                else:
                    formatted_diff.append({
                        'type': 'unchanged',
                        'content': line[2:]
                    })
            
            return Response({
                'version1': {
                    'id': version1.id,
                    'title': version1.title,
                    'created_at': version1.created_at,
                    'user': version1.user.username
                },
                'version2': {
                    'id': version2.id,
                    'title': version2.title,
                    'created_at': version2.created_at,
                    'user': version2.user.username
                },
                'diff': formatted_diff
            })
            
        except NoteVersion.DoesNotExist:
            return Response(
                {'error': 'Version not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['get'])
    def version_history(self, request, pk=None):
        """获取笔记的版本历史"""
        note = self.get_object()
        versions = note.versions.all().order_by('-created_at')
        
        # 计算每个版本的变化
        version_history = []
        prev_version = None
        for version in versions:
            changes = {
                'id': version.id,
                'title': version.title,
                'created_at': version.created_at,
                'user': version.user.username,
                'changes': []
            }
            
            if prev_version:
                # 计算标题变化
                if version.title != prev_version.title:
                    changes['changes'].append({
                        'type': 'title',
                        'old': prev_version.title,
                        'new': version.title
                    })
                
                # 计算内容变化
                content_diff = difflib.unified_diff(
                    prev_version.content.splitlines(),
                    version.content.splitlines(),
                    lineterm=''
                )
                changes['changes'].append({
                    'type': 'content',
                    'diff': list(content_diff)
                })
            
            version_history.append(changes)
            prev_version = version
        
        return Response(version_history)

    @action(detail=True, methods=['post'])
    def restore_version(self, request, pk=None):
        """恢复笔记到指定版本"""
        note = self.get_object()
        version_id = request.data.get('version_id')
        
        if not version_id:
            return Response(
                {'error': 'version_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            version = note.versions.get(id=version_id)
            
            # 保存当前版本到历史记录
            NoteVersion.objects.create(
                note=note,
                title=note.title,
                content=note.content,
                user=request.user
            )
            
            # 恢复版本
            note.title = version.title
            note.content = version.content
            note.save()
            
            # 创建通知
            Notification.objects.create(
                user=request.user,
                title='版本已恢复',
                message=f'笔记 "{note.title}" 已恢复到版本 {version.created_at.strftime("%Y-%m-%d %H:%M")}',
                type='version_restored',
                related_note=note
            )
            
            return Response({'status': 'success'})
            
        except NoteVersion.DoesNotExist:
            return Response(
                {'error': 'Version not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['post'])
    def upload_attachment(self, request, pk=None):
        note = self.get_object()
        file = request.FILES.get('file')
        if not file:
            return Response(
                {'error': 'No file provided'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        # 保存文件
        file_path = f'note_attachments/{note.id}/{file.name}'
        file_name = default_storage.save(file_path, ContentFile(file.read()))

        # 创建附件记录
        attachment = NoteAttachment.objects.create(
            note=note,
            file=file_name,
            file_name=file.name,
            file_size=file.size,
            file_type=file.content_type
        )

        serializer = NoteAttachmentSerializer(attachment)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['get'])
    def attachments(self, request, pk=None):
        note = self.get_object()
        attachments = note.attachments.all()
        serializer = NoteAttachmentSerializer(attachments, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def sync(self, request, pk=None):
        note = self.get_object()
        device_id = request.data.get('device_id')
        version = request.data.get('version', 1)
        
        if not device_id:
            return Response(
                {'error': 'device_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            sync_record = NoteSync.objects.get(note=note, device_id=device_id)
            if version < sync_record.version:
                # 客户端版本较旧，需要更新
                return Response({
                    'status': 'update_required',
                    'note': NoteSerializer(note).data,
                    'version': sync_record.version
                })
            elif version > sync_record.version:
                # 服务器版本较旧，需要更新
                sync_record.version = version
                sync_record.last_synced = timezone.now()
                sync_record.save()
                return Response({'status': 'updated'})
            else:
                # 版本相同，无需更新
                sync_record.last_synced = timezone.now()
                sync_record.save()
                return Response({'status': 'synced'})
        except NoteSync.DoesNotExist:
            # 创建新的同步记录
            NoteSync.objects.create(
                note=note,
                device_id=device_id,
                version=version,
                last_synced=timezone.now()
            )
            return Response({'status': 'synced'})

    @action(detail=False, methods=['get'])
    def sync_status(self, request):
        device_id = request.query_params.get('device_id')
        if not device_id:
            return Response(
                {'error': 'device_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        sync_records = NoteSync.objects.filter(
            note__user=request.user,
            device_id=device_id
        )
        serializer = NoteSyncSerializer(sync_records, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def assign_category(self, request, pk=None):
        note = self.get_object()
        category_id = request.data.get('category_id')
        
        if not category_id:
            return Response(
                {'error': 'category_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            category = Category.objects.get(id=category_id, user=request.user)
            note.category = category
            note.save()
            return Response({'status': 'success'})
        except Category.DoesNotExist:
            return Response(
                {'error': 'Category not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['post'])
    def remove_category(self, request, pk=None):
        note = self.get_object()
        note.category = None
        note.save()
        return Response({'status': 'success'})

    @action(detail=False, methods=['post'])
    def empty_trash(self, request):
        # 获取所有已删除的笔记
        deleted_notes = self.get_queryset().filter(
            is_deleted=True,
            user=request.user
        )
        
        # 删除笔记及其相关数据
        for note in deleted_notes:
            # 删除附件文件
            for attachment in note.attachments.all():
                default_storage.delete(attachment.file.path)
            
            # 删除笔记
            note.delete()
        
        return Response({'status': 'success'})

    @action(detail=True, methods=['get'])
    def export(self, request, pk=None):
        note = self.get_object()
        export_format = request.query_params.get('format', 'markdown')
        
        if export_format == 'markdown':
            content = f"# {note.title}\n\n{note.content}"
            response = HttpResponse(content, content_type='text/markdown')
            response['Content-Disposition'] = f'attachment; filename="{note.title}.md"'
            return response
            
        elif export_format == 'html':
            content = markdown.markdown(note.content)
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>{note.title}</title>
                <meta charset="utf-8">
            </head>
            <body>
                <h1>{note.title}</h1>
                {content}
            </body>
            </html>
            """
            response = HttpResponse(html_content, content_type='text/html')
            response['Content-Disposition'] = f'attachment; filename="{note.title}.html"'
            return response
            
        elif export_format == 'docx':
            document = Document()
            document.add_heading(note.title, 0)
            document.add_paragraph(note.content)
            
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{note.title}.docx"'
            document.save(response)
            return response
            
        elif export_format == 'json':
            data = {
                'title': note.title,
                'content': note.content,
                'created_at': note.created_at.isoformat(),
                'updated_at': note.updated_at.isoformat(),
                'tags': [tag.name for tag in note.tags.all()],
                'category': note.category.name if note.category else None
            }
            response = HttpResponse(json.dumps(data, ensure_ascii=False), content_type='application/json')
            response['Content-Disposition'] = f'attachment; filename="{note.title}.json"'
            return response
            
        else:
            return Response(
                {'error': 'Unsupported export format'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['get'])
    def export_all(self, request):
        notes = self.get_queryset().filter(user=request.user)
        export_format = request.query_params.get('format', 'json')
        
        if export_format == 'json':
            data = []
            for note in notes:
                data.append({
                    'title': note.title,
                    'content': note.content,
                    'created_at': note.created_at.isoformat(),
                    'updated_at': note.updated_at.isoformat(),
                    'tags': [tag.name for tag in note.tags.all()],
                    'category': note.category.name if note.category else None
                })
            response = HttpResponse(json.dumps(data, ensure_ascii=False), content_type='application/json')
            response['Content-Disposition'] = 'attachment; filename="notes_export.json"'
            return response
            
        elif export_format == 'csv':
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = 'attachment; filename="notes_export.csv"'
            
            writer = csv.writer(response)
            writer.writerow(['Title', 'Content', 'Created At', 'Updated At', 'Tags', 'Category'])
            
            for note in notes:
                writer.writerow([
                    note.title,
                    note.content,
                    note.created_at.isoformat(),
                    note.updated_at.isoformat(),
                    ', '.join(tag.name for tag in note.tags.all()),
                    note.category.name if note.category else ''
                ])
            return response
            
        else:
            return Response(
                {'error': 'Unsupported export format'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=True, methods=['post'])
    def add_comment(self, request, pk=None):
        note = self.get_object()
        content = request.data.get('content')
        parent_id = request.data.get('parent_id')
        
        if not content:
            return Response(
                {'error': 'content is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )

        comment = NoteComment.objects.create(
            note=note,
            user=request.user,
            content=content,
            parent_id=parent_id
        )
        serializer = NoteCommentSerializer(comment)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['get'])
    def comments(self, request, pk=None):
        note = self.get_object()
        comments = note.comments.filter(parent=None).order_by('-created_at')
        serializer = NoteCommentSerializer(comments, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def start_collaboration(self, request, pk=None):
        note = self.get_object()
        collaborator_id = request.data.get('collaborator_id')
        role = request.data.get('role', 'editor')
        
        if not collaborator_id:
            return Response(
                {'error': 'collaborator_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            collaborator = User.objects.get(id=collaborator_id)
            collaboration = NoteCollaboration.objects.create(
                note=note,
                collaborator=collaborator,
                role=role,
                invited_by=request.user
            )
            serializer = NoteCollaborationSerializer(collaboration)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        except User.DoesNotExist:
            return Response(
                {'error': 'User not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['get'])
    def collaborators(self, request, pk=None):
        note = self.get_object()
        collaborations = note.collaborations.all()
        serializer = NoteCollaborationSerializer(collaborations, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def update_collaboration(self, request, pk=None):
        note = self.get_object()
        collaboration_id = request.data.get('collaboration_id')
        role = request.data.get('role')
        
        if not collaboration_id or not role:
            return Response(
                {'error': 'collaboration_id and role are required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            collaboration = note.collaborations.get(id=collaboration_id)
            collaboration.role = role
            collaboration.save()
            serializer = NoteCollaborationSerializer(collaboration)
            return Response(serializer.data)
        except NoteCollaboration.DoesNotExist:
            return Response(
                {'error': 'Collaboration not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['post'])
    def save_as_template(self, request, pk=None):
        note = self.get_object()
        template_name = request.data.get('name', note.title)
        
        template = NoteTemplate.objects.create(
            name=template_name,
            content=note.content,
            user=request.user
        )
        serializer = NoteTemplateSerializer(template)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=False, methods=['post'])
    def create_from_template(self, request):
        template_id = request.data.get('template_id')
        title = request.data.get('title')
        
        if not template_id or not title:
            return Response(
                {'error': 'template_id and title are required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            template = NoteTemplate.objects.get(id=template_id, user=request.user)
            note = Note.objects.create(
                title=title,
                content=template.content,
                user=request.user
            )
            serializer = NoteSerializer(note)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        except NoteTemplate.DoesNotExist:
            return Response(
                {'error': 'Template not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=False, methods=['post'])
    def create_backup(self, request):
        """创建笔记备份"""
        try:
            # 获取用户的所有笔记
            notes = Note.objects.filter(user=request.user)
            
            # 创建备份数据
            backup_data = {
                'notes': [],
                'categories': [],
                'tags': [],
                'created_at': timezone.now().isoformat(),
            }
            
            # 添加笔记数据
            for note in notes:
                note_data = {
                    'id': note.id,
                    'title': note.title,
                    'content': note.content,
                    'created_at': note.created_at.isoformat(),
                    'updated_at': note.updated_at.isoformat(),
                    'category_id': note.category_id,
                    'tags': list(note.tags.values_list('id', flat=True)),
                }
                backup_data['notes'].append(note_data)
            
            # 添加分类数据
            categories = Category.objects.filter(user=request.user)
            for category in categories:
                category_data = {
                    'id': category.id,
                    'name': category.name,
                    'created_at': category.created_at.isoformat(),
                }
                backup_data['categories'].append(category_data)
            
            # 添加标签数据
            tags = Tag.objects.filter(user=request.user)
            for tag in tags:
                tag_data = {
                    'id': tag.id,
                    'name': tag.name,
                    'created_at': tag.created_at.isoformat(),
                }
                backup_data['tags'].append(tag_data)
            
            # 创建备份记录
            backup = NoteBackup.objects.create(
                user=request.user,
                name=f'backup_{timezone.now().strftime("%Y%m%d_%H%M%S")}',
                data=backup_data,
            )
            
            # 创建通知
            Notification.objects.create(
                user=request.user,
                title='备份已创建',
                message=f'笔记备份已创建，包含 {len(notes)} 条笔记',
                type='backup_created',
            )
            
            serializer = NoteBackupSerializer(backup)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_400_BAD_REQUEST
            )
    
    @action(detail=True, methods=['post'])
    def restore(self, request, pk=None):
        """从备份恢复笔记"""
        backup = self.get_object()
        
        try:
            with transaction.atomic():
                # 恢复分类
                for category_data in backup.data['categories']:
                    Category.objects.update_or_create(
                        id=category_data['id'],
                        defaults={
                            'name': category_data['name'],
                            'user': request.user,
                        }
                    )
                
                # 恢复标签
                for tag_data in backup.data['tags']:
                    Tag.objects.update_or_create(
                        id=tag_data['id'],
                        defaults={
                            'name': tag_data['name'],
                            'user': request.user,
                        }
                    )
                
                # 恢复笔记
                for note_data in backup.data['notes']:
                    note = Note.objects.update_or_create(
                        id=note_data['id'],
                        defaults={
                            'title': note_data['title'],
                            'content': note_data['content'],
                            'user': request.user,
                            'category_id': note_data['category_id'],
                        }
                    )[0]
                    
                    # 恢复标签关联
                    if note_data['tags']:
                        note.tags.set(note_data['tags'])
                
                # 创建通知
                Notification.objects.create(
                    user=request.user,
                    title='备份已恢复',
                    message=f'已从备份恢复 {len(backup.data["notes"])} 条笔记',
                    type='backup_restored',
                )
                
                return Response({'status': 'success'})
                
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_400_BAD_REQUEST
            )
    
    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """下载备份文件"""
        backup = self.get_object()
        
        try:
            # 创建JSON文件
            response = HttpResponse(
                json.dumps(backup.data, ensure_ascii=False),
                content_type='application/json'
            )
            response['Content-Disposition'] = f'attachment; filename="{backup.name}.json"'
            return response
            
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_400_BAD_REQUEST
            )


class CategoryViewSet(viewsets.ModelViewSet):
    """分类视图集"""
    serializer_class = CategorySerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        return Category.objects.filter(user=self.request.user)
    
    def perform_create(self, serializer):
        serializer.save(user=self.request.user)
    
        # 创建通知
        category = serializer.instance
        Notification.objects.create(
            user=self.request.user,
            title='分类已创建',
            message=f'分类 "{category.name}" 已创建',
            type='category_created'
        )
    
    @action(detail=True, methods=['post'])
    def move_notes(self, request, pk=None):
        """移动笔记到分类"""
        category = self.get_object()
        note_ids = request.data.get('note_ids', [])
        
        if not note_ids:
            return Response(
                {'error': 'note_ids is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            notes = Note.objects.filter(
                id__in=note_ids,
                user=request.user
            )
            
            for note in notes:
                note.category = category
                note.save()
            
            # 创建通知
            Notification.objects.create(
                user=request.user,
                title='笔记已移动',
                message=f'{len(notes)} 条笔记已移动到分类 "{category.name}"',
                type='notes_moved'
            )
            
            return Response({'status': 'success'})
            
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_400_BAD_REQUEST
            )
    
    @action(detail=False, methods=['get'])
    def statistics(self, request):
        """获取分类统计"""
        categories = Category.objects.filter(user=request.user)
        
        statistics = []
        for category in categories:
            note_count = Note.objects.filter(
                category=category,
                user=request.user
            ).count()
            
            statistics.append({
                'category_id': category.id,
                'category_name': category.name,
                'note_count': note_count
            })
        
        return Response(statistics)
    
    @action(detail=False, methods=['get'])
    def tree(self, request):
        """获取分类树"""
        categories = Category.objects.filter(user=request.user)
        
        def build_tree(parent=None):
            tree = []
            for category in categories.filter(parent=parent):
                children = build_tree(category)
                tree.append({
                    'id': category.id,
                    'name': category.name,
                    'children': children
                })
            return tree
        
        return Response(build_tree())
    
    @action(detail=True, methods=['post'])
    def merge(self, request, pk=None):
        """合并分类"""
        category = self.get_object()
        target_category_id = request.data.get('target_category_id')
        
        if not target_category_id:
            return Response(
                {'error': 'target_category_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            target_category = Category.objects.get(
                id=target_category_id,
                user=request.user
            )
            
            # 移动笔记
        notes = Note.objects.filter(
            category=category,
                user=request.user
            )
            notes.update(category=target_category)
            
            # 移动子分类
            subcategories = Category.objects.filter(
                parent=category,
                user=request.user
            )
            subcategories.update(parent=target_category)
            
            # 创建通知
            Notification.objects.create(
            user=request.user,
                title='分类已合并',
                message=f'分类 "{category.name}" 已合并到分类 "{target_category.name}"',
                type='categories_merged'
            )
            
            # 删除原分类
            category.delete()
        except Category.DoesNotExist:
            return Response(
                {'error': 'Target category not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_400_BAD_REQUEST
            )    
            return Response({'status': 'success'})
            



class TagViewSet(viewsets.ModelViewSet):
    """标签视图集"""
    serializer_class = TagSerializer
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [filters.SearchFilter]
    search_fields = ['name']
    
    def get_queryset(self):
        return Tag.objects.filter(user=self.request.user)
    
    def perform_create(self, serializer):
        serializer.save(user=self.request.user)
    
    @action(detail=True, methods=['get'])
    def notes(self, request, pk=None):
        """获取标签下的所有笔记"""
        tag = self.get_object()
        notes = tag.notes.filter(
            user=request.user,
            is_deleted=False
        )
        page = self.paginate_queryset(notes)
        if page is not None:
            serializer = NoteListSerializer(page, many=True)
            return self.get_paginated_response(serializer.data)
        serializer = NoteListSerializer(notes, many=True)
        return Response(serializer.data)


class NoteShareViewSet(viewsets.ModelViewSet):
    """笔记分享视图集"""
    serializer_class = NoteShareSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return NoteVersion.objects.filter(note__user=self.request.user)
    
    def perform_create(self, serializer):
        serializer.save()
        
        # 创建通知
        version = serializer.instance
        Notification.objects.create(
            user=self.request.user,
            title='新版本已创建',
            message=f'笔记 "{version.note.title}" 的新版本已创建',
            type='version_created',
            related_note=version.note
        )

    @action(detail=True, methods=['post'])
    def restore(self, request, pk=None):
        """恢复版本"""
        version = self.get_object()
        
        try:
            # 保存当前版本
            current_version = NoteVersion.objects.create(
                note=version.note,
                title=version.note.title,
                content=version.note.content,
                user=request.user
            )
            
            # 恢复版本
            version.note.title = version.title
            version.note.content = version.content
            version.note.save()
            
            # 创建通知
            Notification.objects.create(
                user=request.user,
                title='版本已恢复',
                message=f'笔记 "{version.note.title}" 已恢复到版本 {version.version_number}',
                type='version_restored',
                related_note=version.note
            )
        
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_400_BAD_REQUEST
            )    
        return Response({'status': 'success'})


    
    @action(detail=False, methods=['get'])
    def history(self, request):
        """获取版本历史"""
        note_id = request.query_params.get('note_id')
        
        if not note_id:
            return Response(
                {'error': 'note_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        versions = NoteVersion.objects.filter(
            note_id=note_id,
            note__user=request.user
        ).order_by('-created_at')
        
        serializer = NoteVersionSerializer(versions, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'])
    def diff(self, request, pk=None):
        """获取版本差异"""
        version = self.get_object()
        compare_version_id = request.query_params.get('compare_version_id')
        
        if not compare_version_id:
            return Response(
                {'error': 'compare_version_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            compare_version = NoteVersion.objects.get(
                id=compare_version_id,
                note__user=request.user
            )
            
            # 计算差异
            differ = difflib.Differ()
            diff = list(differ.compare(
                compare_version.content.splitlines(),
                version.content.splitlines()
            ))
            
            return Response({
                'diff': diff,
                'current_version': version.version_number,
                'compare_version': compare_version.version_number
            })
            
        except NoteVersion.DoesNotExist:
            return Response(
                {'error': 'Compare version not found'}, 
                status=status.HTTP_404_NOT_FOUND
            )
    
    @action(detail=False, methods=['get'])
    def statistics(self, request):
        """获取版本统计"""
        note_id = request.query_params.get('note_id')
        
        if not note_id:
            return Response(
                {'error': 'note_id is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        versions = NoteVersion.objects.filter(
            note_id=note_id,
            note__user=request.user
        )
        
        statistics = {
            'total_versions': versions.count(),
            'first_version': versions.order_by('created_at').first().created_at if versions.exists() else None,
            'last_version': versions.order_by('-created_at').first().created_at if versions.exists() else None,
            'versions_per_month': versions.annotate(
                month=TruncMonth('created_at')
            ).values('month').annotate(
                count=Count('id')
            ).order_by('month')
        }
        
        return Response(statistics)