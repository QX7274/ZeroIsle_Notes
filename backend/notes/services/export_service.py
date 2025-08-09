from django.core.files.base import ContentFile
import docx
from io import BytesIO
from PyPDF2 import PdfWriter
from ..models import Note

class ExportService:
    """
    笔记导出服务，支持PDF和Word格式
    """
    @staticmethod
    def export_to_word(note: Note) -> ContentFile:
        """
        将笔记导出为Word文档
        :param note: 要导出的笔记对象
        :return: 包含Word文档的ContentFile对象
        """
        # 创建Word文档
        doc = docx.Document()
        doc.add_heading(note.title, level=1)
        doc.add_paragraph(note.content)

        # 添加标签
        if note.tags.exists():
            tags = ", ".join([tag.name for tag in note.tags.all()])
            doc.add_paragraph(f"标签: {tags}", style='Intense Quote')

        # 添加元数据
        doc.add_paragraph(f"创建时间: {note.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"更新时间: {note.updated_at.strftime('%Y-%m-%d %H:%M:%S')}")

        # 保存到内存中
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return ContentFile(buffer.getvalue(), f"{note.title}.docx")

    @staticmethod
    def export_to_pdf(note: Note) -> ContentFile:
        """
        将笔记导出为PDF文档
        :param note: 要导出的笔记对象
        :return: 包含PDF文档的ContentFile对象
        """
        # 创建PDF文档
        buffer = BytesIO()
        pdf_writer = PdfWriter()

        # 这里需要添加PDF内容生成逻辑
        # 实际实现中可能需要使用ReportLab或其他PDF生成库
        # 为简化示例，我们创建一个基本PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        pdf_canvas = canvas.Canvas(buffer, pagesize=letter)
        pdf_canvas.setFont("Helvetica-Bold", 16)
        pdf_canvas.drawString(100, 750, note.title)

        pdf_canvas.setFont("Helvetica", 12)
        text_object = pdf_canvas.beginText(100, 720)
        text_object.textLines(note.content)
        pdf_canvas.drawText(text_object)

        pdf_canvas.save()
        buffer.seek(0)

        return ContentFile(buffer.getvalue(), f"{note.title}.pdf")

    @staticmethod
    def export_notes(notes, export_format: str, zip_filename: str = None) -> ContentFile:
        """
        批量导出笔记
        :param notes: 笔记对象列表
        :param export_format: 导出格式，支持 'pdf' 和 'docx'
        :param zip_filename: 压缩包文件名，如果提供则将所有笔记压缩成一个ZIP文件
        :return: 包含导出文件的ContentFile对象
        """
        if export_format not in ['pdf', 'docx']:
            raise ValueError("不支持的导出格式，仅支持pdf和docx")

        if len(notes) == 1 and not zip_filename:
            # 单个笔记直接导出
            export_method = ExportService.export_to_pdf if export_format == 'pdf' else ExportService.export_to_word
            return export_method(notes[0])

        # 多个笔记，需要创建ZIP文件
        import zipfile
        buffer = BytesIO()

        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for note in notes:
                export_method = ExportService.export_to_pdf if export_format == 'pdf' else ExportService.export_to_word
                file_content = export_method(note)
                zipf.writestr(file_content.name, file_content.read())

        buffer.seek(0)
        zip_filename = zip_filename or f"notes_export_{len(notes)}.zip"
        return ContentFile(buffer.getvalue(), zip_filename)